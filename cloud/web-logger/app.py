from flask import Flask, render_template, send_file, redirect
from flask import *
from flask_socketio import SocketIO, emit
from azure.storage.blob import BlobSasPermissions, BlobServiceClient
from azure.cosmos import cosmos_client
from docx import Document
from docx.shared import Inches
from io import BytesIO
import time, threading, requests

### CREDENTIALS ###
COSMOS_ENDPOINT = ""
COSMOS_KEY = ""
COSMOS_DATABASE_NAME = ""
COSMOS_CONTAINER_NAME = ""
STORAGE_ACCOUNT_NAME = ""
STORAGE_CONTAINER_NAME = ""
STORAGE_CONTAINER_CAPTIONED_NAME = ""
STORAGE_KEY = ""
STORAGE_CONNECTION_STRING = ""

# set up clients
app = Flask(__name__)
sio = SocketIO(app)  # socket
client = cosmos_client.CosmosClient(COSMOS_ENDPOINT, COSMOS_KEY)
database = client.get_database_client(COSMOS_DATABASE_NAME)
container = database.get_container_client(COSMOS_CONTAINER_NAME)
blob_service_client = BlobServiceClient.from_connection_string(conn_str=STORAGE_CONNECTION_STRING)

# global variables
global record_new
global record_last
record_new = list(container.query_items("SELECT TOP 1 * FROM c ORDER BY c._ts DESC", enable_cross_partition_query=True))
record_last = record_new

# poll cosmos db every second and update logs
def cosmos_db_polling():
    global record_new  
    while True:
        # get the latest 1 records from cosmos db
        record_new = list(container.query_items("SELECT TOP 1 * FROM c ORDER BY c._ts DESC", enable_cross_partition_query=True))
        time.sleep(3)

# make a new thread 
thread = threading.Thread(target=cosmos_db_polling)
thread.start()

# call back when ping event is invoked
@sio.on("ping") 
def ping():
    global record_new
    global record_last
    # if the latest record is the same as the last record, emit "noupdate"
    if record_new == record_last:
        emit("noupdate")    
    # if the latest record is different from the last record, emit "update"
    else:
        # convert unixtime to human readable time in japan time
        time_readable = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(int(record_new[0]["time"]) + 32400))
        emit("update", {"time": time_readable, "description": record_new[0]["description"], "level": record_new[0]["level"], "measurement": record_new[0]["measurement"], 
                        "reason": record_new[0]["reason"], "image_url": record_new[0]["image_url"], "image_captioned_url": record_new[0]["image_captioned_url"]})
        record_last = record_new

@app.route("/report", methods=["POST"])
def report():
    """Generate a summary report of the logs whose mesaurement is not 'n/a'"""
    records = list(container.query_items("SELECT * FROM c WHERE c.measurement != 'n/a' ORDER BY c._ts DESC", enable_cross_partition_query=True))
    # make a summary report using word docs
    doc = Document()
    doc.add_heading('Summary Report on Potental Dangers', 0)
    for record in records:
        doc.add_heading(record["description"], level=1)
        doc.add_paragraph(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(int(record['time']) + 32400))}")
        doc.add_paragraph(f"Level: {record['level']}")
        doc.add_paragraph(f"Measurement: {record['measurement']}")
        doc.add_paragraph(f"Reason: {record['reason']}")
        # insert the image from the url wwithout saving it to the local disk
        table = doc.add_table(rows=1, cols=2)
        response1 = requests.get(record["image_url"])
        image1 = BytesIO(response1.content)
        response2 = requests.get(record["image_captioned_url"])
        image2 = BytesIO(response2.content)
        # Add the images to the cells
        cell_1 = table.cell(0, 0)
        cell_1.width = Inches(3)
        cell_1.paragraphs[0].add_run().add_picture(image1, width=Inches(3))
        cell_2 = table.cell(0, 1)
        cell_2.width = Inches(3)
        cell_2.paragraphs[0].add_run().add_picture(image2, width=Inches(3))
        doc.add_paragraph("---------------------------------------------------")
    doc.save("summary_report.docx")
    return send_file("summary_report.docx", as_attachment=True)

# delete all items in azure storage blob and all records in cosmos db container
@app.route("/reset", methods=["POST"])
def reset():
    container_client = blob_service_client.get_container_client(container=STORAGE_CONTAINER_NAME)
    delete_blobs(container_client) # delete all blobs in the container
    container_client = blob_service_client.get_container_client(container=STORAGE_CONTAINER_CAPTIONED_NAME)
    delete_blobs(container_client) # delete all blobs in the container
    # detele all items in the cosmos container
    items = container.query_items(query="SELECT * FROM c", enable_cross_partition_query=True)
    for item in items:
        container.delete_item(item, partition_key=item["id"])
    print("*Reset complete!")
    return redirect('/')

def delete_blobs(container_client):
    blobs = container_client.list_blobs()
    blobs_list = list(blobs) # convert ItemPaged object to list
    for batch in chunks(blobs_list, 100): # delete blobs in batches of 100
        container_client.delete_blobs(*batch)

def chunks(lst, n):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), n):
        yield lst[i:i + n]

@app.route("/")
def index():
    return render_template("index.html")
    
if __name__ == "__main__":
    sio.run(app, host='127.0.0.1', port=5000, debug=True) # app.run(host='127.0.0.1', port=5000, debug=True)